from docx import Document
import sys
from pathlib import Path

# Add project root to path
# inspect_template.py lives under scripts/utils/, so go up 2 levels to repo root (apps/narrative/).
BASE_DIR = Path(__file__).resolve().parents[2]
template_path = BASE_DIR / "templates/cgp_p3_template.docx"

if not template_path.exists():
    print(f"Template not found at {template_path}")
    sys.exit(1)

doc = Document(template_path)

print("\n--- PLACEHOLDERS FOUND ---")
count = 0
# Check paragraphs
for i, para in enumerate(doc.paragraphs):
    if "INSERT" in para.text:
        print(f"Para {i}: {para.text.strip()[:100]}...")
        count += 1

# Check tables
for i, table in enumerate(doc.tables):
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            if "INSERT" in cell.text:
                print(f"Table {i+1}, Row {r}, Cell {c}: {cell.text.strip()[:100]}...")
                count += 1

print(f"\nTotal placeholders found: {count}")
